"""
测试报告生成 - Word + 通知（企业微信/飞书/钉钉）
被引用方：09-报告生成 agent
"""
import json
import logging
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional

import requests

logger = logging.getLogger(__name__)


# ===== Word 报告 =====

def generate_test_report(data: Dict, output_path: str) -> str:
    """生成 Word 测试报告。data 包含 project_name/version/environment/results/bugs/coverage/risks 等字段。"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 字体 fallback：微软雅黑 → PingFang SC → Noto Sans CJK SC
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    title = doc.add_heading("测试报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"项目：{data.get('project_name', '')}")
    doc.add_paragraph(f"版本：{data.get('version', '')}")
    doc.add_paragraph(f"测试环境：{data.get('environment', '')}")
    doc.add_paragraph(f"报告日期：{datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_page_break()

    # 一、执行摘要
    doc.add_heading("一、执行摘要", level=1)
    summary_table = doc.add_table(rows=2, cols=5)
    summary_table.style = "Table Grid"
    headers = ["总用例", "通过", "失败", "通过率", "覆盖率"]
    for i, h in enumerate(headers):
        cell = summary_table.cell(0, i)
        cell.text = h
        if cell.paragraphs and cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    results = data.get("results", {})
    values = [
        str(results.get("total", 0)),
        str(results.get("passed", 0)),
        str(results.get("failed", 0)),
        f"{results.get('pass_rate', 0):.1%}",
        f"{data.get('coverage', 0):.1%}",
    ]
    for i, v in enumerate(values):
        summary_table.cell(1, i).text = v

    conclusion = doc.add_paragraph()
    conclusion.add_run("测试结论：").bold = True
    verdict = data.get("verdict", "通过")
    run = conclusion.add_run(verdict)
    run.font.color.rgb = RGBColor(0, 128, 0) if verdict == "通过" else RGBColor(255, 0, 0)
    run.bold = True

    # 二、缺陷统计
    doc.add_heading("二、缺陷统计", level=1)
    bugs = data.get("bugs", {})
    bug_table = doc.add_table(rows=5, cols=3)
    bug_table.style = "Table Grid"
    rows = [
        ("级别", "数量", "修复率"),
        ("P0", str(bugs.get("p0", 0)), f"{bugs.get('p0_fix_rate', 0):.0%}"),
        ("P1", str(bugs.get("p1", 0)), f"{bugs.get('p1_fix_rate', 0):.0%}"),
        ("P2", str(bugs.get("p2", 0)), f"{bugs.get('p2_fix_rate', 0):.0%}"),
        ("P3", str(bugs.get("p3", 0)), f"{bugs.get('p3_fix_rate', 0):.0%}"),
    ]
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = bug_table.cell(ri, ci)
            cell.text = val
            if ri == 0 and cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True

    # 三、性能指标（可选）
    perf = data.get("performance")
    if perf:
        doc.add_heading("三、性能指标（JMeter）", level=1)
        p_table = doc.add_table(rows=5, cols=2)
        p_table.style = "Table Grid"
        p_rows = [
            ("TPS", f"{perf.get('tps', 0)} 次/秒"),
            ("平均响应", f"{perf.get('avg_response_ms', 0)} ms"),
            ("P95响应", f"{perf.get('p95_response_ms', 0)} ms"),
            ("错误率", f"{perf.get('error_rate_pct', 0)}%"),
            ("门禁结论", perf.get("quality_gate", "PASS")),
        ]
        for ri, (k, v) in enumerate(p_rows):
            p_table.cell(ri, 0).text = k
            p_table.cell(ri, 1).text = v

    # 四、风险与建议
    doc.add_heading("四、风险与建议", level=1)
    for risk in data.get("risks", []):
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(f"【{risk.get('level', '中')}】").bold = True
        para.add_run(risk.get("description", ""))

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"Word 报告已生成: {output_path}")
    return output_path


# ===== 通知 =====

def send_wechat_report(summary: Dict, webhook: Optional[str] = None) -> bool:
    """企业微信 markdown 通知"""
    webhook = webhook or os.getenv("WECHAT_WEBHOOK_URL")
    if not webhook:
        logger.warning("未配置 WECHAT_WEBHOOK_URL，跳过企业微信通知")
        return False
    pass_rate = summary.get("pass_rate", 0)
    verdict = summary.get("verdict", "通过")
    icon = "✅" if verdict == "通过" else "❌"
    content = (
        f"{icon} **{summary.get('project','')} 测试完成**\n"
        f"> 环境：{summary.get('environment','')}\n"
        f"> 用例：{summary.get('total',0)} 通过 {summary.get('passed',0)}（{pass_rate:.1%}）\n"
        f"> Bug：P0({summary.get('p0_bugs',0)}) P1({summary.get('p1_bugs',0)})\n"
        f"> 性能：TPS={summary.get('perf_tps','-')} P95={summary.get('perf_p95','-')}ms\n"
        f"> [查看完整报告]({summary.get('report_url','#')})"
    )
    try:
        r = requests.post(webhook, json={"msgtype": "markdown", "markdown": {"content": content}}, timeout=10)
        return r.ok
    except Exception as e:
        logger.error(f"企业微信通知失败: {e}")
        return False


def send_feishu_report(summary: Dict, webhook: Optional[str] = None) -> bool:
    """飞书 富文本卡片 通知"""
    webhook = webhook or os.getenv("FEISHU_WEBHOOK")
    if not webhook:
        logger.warning("未配置 FEISHU_WEBHOOK，跳过飞书通知")
        return False
    # 飞书卡片合法颜色：blue/wathet/turquoise/green/yellow/orange/red/carmine/violet/purple/indigo/grey
    color = "green" if summary.get("pass_rate", 0) >= 0.9 else "red"
    card = {
        "config": {"wide_screen_mode": True},
        "header": {
            "title": {"content": f"测试报告 - {summary.get('project','')}", "tag": "plain_text"},
            "template": color,
        },
        "elements": [
            {
                "tag": "div",
                "fields": [
                    {"is_short": True, "text": {"tag": "lark_md", "content": f"**环境**\n{summary.get('environment','')}"}},
                    {"is_short": True, "text": {"tag": "lark_md", "content": f"**通过率**\n{summary.get('pass_rate',0):.1%}"}},
                    {"is_short": True, "text": {"tag": "lark_md", "content": f"**P0 Bug**\n{summary.get('p0_bugs',0)}个"}},
                    {"is_short": True, "text": {"tag": "lark_md", "content": f"**覆盖率**\n{summary.get('coverage',0):.1%}"}},
                ],
            },
            {
                "tag": "action",
                "actions": [{
                    "tag": "button",
                    "text": {"content": "查看完整报告", "tag": "plain_text"},
                    "url": summary.get("report_url", "#"),
                    "type": "primary",
                }],
            },
        ],
    }
    try:
        r = requests.post(webhook, json={"msg_type": "interactive", "card": card}, timeout=10)
        return r.ok
    except Exception as e:
        logger.error(f"飞书通知失败: {e}")
        return False


def send_dingtalk_report(summary: Dict, webhook: Optional[str] = None) -> bool:
    """钉钉 markdown 通知"""
    webhook = webhook or os.getenv("DINGTALK_WEBHOOK")
    if not webhook:
        logger.warning("未配置 DINGTALK_WEBHOOK，跳过钉钉通知")
        return False
    pass_rate = summary.get("pass_rate", 0)
    verdict = summary.get("verdict", "通过")
    icon = "✅" if verdict == "通过" else "❌"
    text = (
        f"### {icon} {summary.get('project','')} 测试完成\n"
        f"- 环境：{summary.get('environment','')}\n"
        f"- 通过率：{pass_rate:.1%}\n"
        f"- Bug：P0({summary.get('p0_bugs',0)}) P1({summary.get('p1_bugs',0)})\n"
        f"- 性能：TPS={summary.get('perf_tps','-')} P95={summary.get('perf_p95','-')}ms\n"
        f"- [查看报告]({summary.get('report_url','#')})"
    )
    try:
        r = requests.post(webhook, json={
            "msgtype": "markdown",
            "markdown": {"title": "测试报告", "text": text},
        }, timeout=10)
        return r.ok
    except Exception as e:
        logger.error(f"钉钉通知失败: {e}")
        return False


def send_all_notifications(summary: Dict):
    """同时发送企业微信/飞书/钉钉（按 .env 配置自动跳过未配置项）"""
    return {
        "wechat": send_wechat_report(summary),
        "feishu": send_feishu_report(summary),
        "dingtalk": send_dingtalk_report(summary),
    }


# ===== CLI =====

def main():
    import argparse
    logging.basicConfig(level=logging.INFO)
    parser = argparse.ArgumentParser(description="生成 Word 测试报告 + 通知")
    parser.add_argument("--data", required=True, help="测试结果 JSON 文件路径")
    parser.add_argument("--output", default=None, help="Word 报告输出路径")
    parser.add_argument("--notify", action="store_true", help="发送通知")
    args = parser.parse_args()

    with open(args.data, encoding="utf-8") as f:
        data = json.load(f)

    output = args.output or f"workspace/执行日志/报告/测试报告_{datetime.now().strftime('%Y%m%d')}.docx"
    generate_test_report(data, output)

    if args.notify:
        results = send_all_notifications(data)
        print(f"通知发送结果：{results}")


if __name__ == "__main__":
    main()
